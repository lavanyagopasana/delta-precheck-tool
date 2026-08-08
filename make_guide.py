# -*- coding: utf-8 -*-
"""
Generates a non-technical Word (.docx) user guide for the Delta Migration Readiness Tracker,
with brand-styled flow diagrams. Process/direction focused -- no backend code or logic.
"""

import os
import tempfile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Circle, Polygon

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- brand palette
PURPLE = "#4B1FA6"
PURPLE_DK = "#3A1580"
PURPLE_LT = "#EDE7FB"
BLUE = "#0129AC"
GREEN = "#16A34A"
GREEN_LT = "#E5F8ED"
YELLOW = "#B45309"
YELLOW_LT = "#FEF6E0"
RED = "#DC2626"
INK = "#2E2E2E"
MUTED = "#6B7385"
TRACK = "#E7EAF0"

ASSET_DIR = tempfile.mkdtemp(prefix="delta_guide_")


def _hex(c):
    return c.lstrip("#")


# ================================================================ DIAGRAMS
def _box(ax, x, y, w, h, title, subtitle=None, fill=PURPLE_LT, edge=PURPLE,
         tcolor=INK, number=None, title_size=12, sub_size=9.5, bold=True):
    ax.add_patch(FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.01,rounding_size=0.10",
        linewidth=1.4, edgecolor=edge, facecolor=fill, zorder=2))
    tx = x + 0.28
    if number is not None:
        cx, cy = x + 0.34, y + h / 2
        ax.add_patch(Circle((cx, cy), 0.24, facecolor=PURPLE, edgecolor="none", zorder=3))
        ax.text(cx, cy, str(number), ha="center", va="center", color="white",
                fontsize=11, fontweight="bold", zorder=4)
        tx = x + 0.78
    if subtitle:
        ax.text(tx, y + h * 0.62, title, ha="left", va="center", color=tcolor,
                fontsize=title_size, fontweight="bold" if bold else "normal", zorder=4)
        ax.text(tx, y + h * 0.30, subtitle, ha="left", va="center", color=MUTED,
                fontsize=sub_size, zorder=4)
    else:
        ax.text(tx, y + h / 2, title, ha="left", va="center", color=tcolor,
                fontsize=title_size, fontweight="bold" if bold else "normal", zorder=4)


def _arrow(ax, x1, y1, x2, y2, color=PURPLE, style="-|>", lw=1.8, rad=0.0):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2), arrowstyle=style, mutation_scale=16,
        linewidth=lw, color=color,
        connectionstyle=f"arc3,rad={rad}", zorder=1))


def diagram_lifecycle(path):
    steps = [
        ("Set up the project & its servers", "A project is created; each server to migrate is added"),
        ("Import the workspace pairs", "A CSV lists every source \u2192 destination account for that server"),
        ("Complete the pre-checks", "Each readiness item gets a status, a note, and evidence attached"),
        ("Submit the pre-check", "Once every item is done, the engineer submits it for review"),
        ("Approvals sign-off chain", "Migration Manager \u2192 Dev Lead \u2192 QA Lead review in order"),
        ("Delta Ready", "All approvals granted \u2014 the server is cleared to migrate"),
        ("Initiate Delta (Start)", "The engineer starts the actual data migration"),
        ("Finish Delta", "Migration is completed and recorded"),
    ]
    n = len(steps)
    fig, ax = plt.subplots(figsize=(7.4, 10.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, n * 1.28 + 0.3)
    ax.axis("off")
    bw, bh = 8.7, 0.95
    x = 0.65
    for i, (t, s) in enumerate(steps):
        y = (n - 1 - i) * 1.28 + 0.25
        last = (i == n - 1)
        ready = (i == 5)
        fill = GREEN_LT if (ready or last) else PURPLE_LT
        edge = GREEN if (ready or last) else PURPLE
        _box(ax, x, y, bw, bh, t, s, fill=fill, edge=edge, number=i + 1)
        if i < n - 1:
            _arrow(ax, x + bw / 2, y, x + bw / 2, y - 0.33, color=PURPLE)
    ax.text(5.0, n * 1.28 + 0.12, "The journey of every server",
            ha="center", va="center", fontsize=13.5, fontweight="bold", color=PURPLE_DK)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def diagram_approval(path):
    fig, ax = plt.subplots(figsize=(8.6, 8.2))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 11)
    ax.axis("off")
    bw, bh = 6.4, 0.9
    cx = 1.4

    def cbox(y, t, s, fill=PURPLE_LT, edge=PURPLE):
        _box(ax, cx, y, bw, bh, t, s, fill=fill, edge=edge)

    cbox(9.6, "Pre-check submitted", "Engineer finishes and submits the server's checklist")
    _arrow(ax, cx + bw / 2, 9.6, cx + bw / 2, 9.25)
    cbox(8.3, "Migration Manager reviews", "Approves that the pre-checks are complete and valid")
    _arrow(ax, cx + bw / 2, 8.3, cx + bw / 2, 7.95)
    cbox(7.0, "Dev Lead reviews", "Approves \u2014 and decides whether QA is needed")
    _arrow(ax, cx + bw / 2, 7.0, cx + bw / 2, 6.5)

    # decision diamond
    dcx, dcy = cx + bw / 2, 5.7
    dw = 1.5
    ax.add_patch(Polygon([(dcx, dcy + 0.8), (dcx + dw, dcy), (dcx, dcy - 0.8), (dcx - dw, dcy)],
                         closed=True, facecolor=YELLOW_LT, edgecolor=YELLOW, linewidth=1.4, zorder=2))
    ax.text(dcx, dcy, "QA needed?", ha="center", va="center", fontsize=11,
            fontweight="bold", color=INK, zorder=4)

    # Yes -> QA lead
    _arrow(ax, dcx, dcy - 0.8, dcx, 4.05)
    ax.text(dcx + 0.25, 4.55, "Yes", fontsize=10, color=YELLOW, fontweight="bold")
    cbox(3.15, "QA Lead reviews", "Final approval on migration quality")
    _arrow(ax, cx + bw / 2, 3.15, cx + bw / 2, 2.55)

    # No -> straight to ready (curve right)
    _arrow(ax, dcx + dw, dcy, 11.2, dcy, rad=0.0)
    ax.text(dcx + dw + 0.4, dcy + 0.28, "No (QA skipped)", fontsize=9.5, color=MUTED)
    _arrow(ax, 11.2, dcy, 11.2, 2.15, rad=0.0, style="-")
    _arrow(ax, 11.2, 2.15, cx + bw / 2 + 0.2, 2.15, rad=0.0)

    cbox(1.25, "Delta Ready", "Every required gate is green \u2014 the server can migrate",
         fill=GREEN_LT, edge=GREEN)

    # decline note
    _arrow(ax, cx, 7.45, cx - 0.9, 8.75, color=RED, rad=-0.35, style="-|>")
    ax.text(cx - 1.0, 8.05, "Declined?\nreturns to the\nprevious step",
            fontsize=8.5, color=RED, ha="center", va="center")

    ax.text(6.0, 10.6, "How approvals flow", ha="center", va="center",
            fontsize=13.5, fontweight="bold", color=PURPLE_DK)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def diagram_notifications(path):
    events = [
        ("Pre-check\nsubmitted", "Migration Manager"),
        ("Manager\napproves", "Dev Lead"),
        ("QA required\n(Dev decides)", "QA Lead"),
        ("All approvals\ngranted", "Migration Manager\n(Delta Ready)"),
        ("Engineer clicks\nStart", "Migration Manager\n(Delta Initiated)"),
        ("Engineer clicks\nFinish", "Migration Manager\n(Delta Finished)"),
    ]
    n = len(events)
    fig, ax = plt.subplots(figsize=(9.6, 4.4))
    ax.set_xlim(0, n * 2.0)
    ax.set_ylim(0, 5)
    ax.axis("off")
    bw = 1.7
    for i, (ev, who) in enumerate(events):
        x = i * 2.0 + 0.12
        ax.add_patch(FancyBboxPatch((x, 3.1), bw, 1.1, boxstyle="round,pad=0.02,rounding_size=0.12",
                                    facecolor=PURPLE_LT, edgecolor=PURPLE, linewidth=1.3, zorder=2))
        ax.text(x + bw / 2, 3.65, ev, ha="center", va="center", fontsize=9, color=INK, fontweight="bold")
        # down arrow
        _arrow(ax, x + bw / 2, 3.1, x + bw / 2, 2.35, color=GREEN)
        ax.text(x + bw / 2, 2.55, "email", ha="center", va="bottom", fontsize=7.5, color=GREEN)
        ax.add_patch(FancyBboxPatch((x, 1.05), bw, 1.15, boxstyle="round,pad=0.02,rounding_size=0.12",
                                    facecolor=GREEN_LT, edgecolor=GREEN, linewidth=1.3, zorder=2))
        ax.text(x + bw / 2, 1.62, who, ha="center", va="center", fontsize=8.3, color=INK)
        if i < n - 1:
            _arrow(ax, x + bw, 3.65, x + 2.0 + 0.12, 3.65, color=MUTED, lw=1.4)
    ax.text(n * 2.0 / 2, 4.7, "Who gets notified, and when", ha="center", va="center",
            fontsize=13, fontweight="bold", color=PURPLE_DK)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ================================================================ DOCX HELPERS
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(hex_color))
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def style_runs(cell, size=10, color=INK, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor.from_string(_hex(color))
            r.font.bold = bold


def add_heading(doc, text, size=16, color=PURPLE, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(_hex(color))
    return p


def add_para(doc, text, size=11, color=INK, italic=False, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(_hex(color))
    return p


def add_bullet(doc, text, size=11, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(_hex(color))
    return p


def add_table(doc, headers, rows, header_bg=PURPLE, widths=None, header_color="#FFFFFF"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        set_cell_margins(hdr[i])
        style_runs(hdr[i], size=10.5, color=header_color, bold=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            set_cell_margins(cells[ci])
            set_cell_bg(cells[ci], "#FFFFFF" if ri % 2 == 0 else "#F6F4FC")
            style_runs(cells[ci], size=10, color=INK, bold=(ci == 0 and len(headers) > 1))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def add_image_centered(doc, path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _hex(PURPLE_LT))
    pbdr.append(bottom)
    pPr.append(pbdr)


# ================================================================ BUILD
def build():
    life = os.path.join(ASSET_DIR, "lifecycle.png")
    appr = os.path.join(ASSET_DIR, "approval.png")
    noti = os.path.join(ASSET_DIR, "notifications.png")
    diagram_lifecycle(life)
    diagram_approval(appr)
    diagram_notifications(noti)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(_hex(INK))

    # ---- Cover
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(40)
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("CLOUDFUZE  \u2022  INTERNAL TOOL  \u2022  USER GUIDE")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(_hex(PURPLE))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Delta Migration Readiness Tracker")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(_hex(PURPLE_DK))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A plain-language guide to what the tool does and how work flows through it")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(_hex(MUTED))
    add_para(doc,
             "This guide explains the tool as a process \u2014 no technical background needed. "
             "It walks through the problem it solves, who uses it, and the exact path every server "
             "takes from first setup to a finished migration.",
             size=11.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    doc.add_page_break()

    # ---- 1. What is this tool
    add_heading(doc, "1. What is this tool?")
    add_para(doc,
             "The Delta Migration Readiness Tracker is one central place for the Migration team to "
             "track and prove that a server is truly ready before its \u201cDelta\u201d migration is "
             "started. \u201cDelta\u201d simply means moving the actual data (mailboxes, drives, or "
             "messages) from a source system to a destination system.")
    add_para(doc,
             "Instead of chasing status across spreadsheets, chats, and emails, everyone looks at the "
             "same screen. A server can only move forward when the right checks are done and the right "
             "people have signed off. Nothing important can be skipped or forgotten.")

    # ---- 2. Problem it solves
    add_heading(doc, "2. The problem it solves")
    add_table(doc, ["Before this tool", "With this tool"], [
        ["Readiness checks, evidence, and approvals were scattered across spreadsheets and email.",
         "Everything lives in one place, always up to date."],
        ["No trustworthy, shared view of what was actually ready.",
         "A single dashboard shows readiness, approvals, and blockers at a glance."],
        ["Easy to start a migration before it was truly signed off.",
         "The migration button stays locked until every required gate is green."],
        ["Hard to see who is waiting on whom.",
         "Each step has a clear owner and an automatic email reminder."],
    ], widths=[3.1, 3.1])

    # ---- 3. How it helps
    add_heading(doc, "3. How it helps the team")
    for b in [
        "One dashboard for server readiness, pending approvals, open tickets, and progress per project.",
        "A guaranteed process \u2014 a Delta cannot start until pre-checks are complete and every required role has approved.",
        "Evidence is captured for every check, so readiness can be proven later, not just claimed.",
        "Automatic email notifications keep the right person informed at every stage.",
        "Clear roles and permissions \u2014 people only see and do what fits their job.",
    ]:
        add_bullet(doc, b)

    # ---- 4. Glossary
    add_heading(doc, "4. Key words in plain English")
    add_table(doc, ["Term", "What it means"], [
        ["Project", "A customer engagement or piece of work. It groups everything below it."],
        ["Server", "A system being migrated. Each project can have several servers."],
        ["Workspace pair", "One source account \u2192 one destination account (e.g., an old mailbox mapped to a new one). Imported from a CSV file."],
        ["Pre-check", "A readiness checklist for a server. Every item needs a status, a note, and an attached evidence file."],
        ["Ticket", "A logged blocker or issue linked to a project/server, marked Open or Resolved."],
        ["Sign-off / Approval", "A required review by a specific role before the work can continue."],
        ["Delta", "The actual data migration from source to destination."],
        ["Decommission-ready", "Every server in a project has finished its Delta \u2014 the whole project is done."],
    ], widths=[1.7, 4.5])

    # ---- 5. Roles
    add_heading(doc, "5. Who uses it \u2014 the roles")
    add_para(doc, "Anyone signs in with their Microsoft account, but what they can do depends on their role:")
    add_table(doc, ["Role", "What they do"], [
        ["Admin", "Manages people and access; can see and do everything."],
        ["Migration Manager", "Owns the project; first to approve after a pre-check is submitted."],
        ["Dev Lead", "Second approver; also decides whether a QA review is required."],
        ["QA Lead", "Final approver on quality, when the Dev Lead requires it."],
        ["Migration Engineer", "Does the hands-on work: imports pairs, completes pre-checks, and starts/finishes the Delta."],
    ], widths=[1.9, 4.3])

    # ---- 6. Lifecycle
    doc.add_page_break()
    add_heading(doc, "6. The journey of a server, step by step")
    add_para(doc,
             "Every server follows the same path. Each stage must be finished before the next one opens \u2014 "
             "this is what guarantees nothing is skipped.")
    add_image_centered(doc, life, 5.5)
    steps_detail = [
        ("Set up the project & its servers", "A Migration Manager or Admin creates the project and adds each server that needs migrating."),
        ("Import the workspace pairs", "An engineer uploads a CSV listing every source \u2192 destination account. Duplicate rows are skipped automatically and reported."),
        ("Complete the pre-checks", "The engineer works through the readiness checklist, adding a status, a short note, and an evidence file to each item."),
        ("Submit the pre-check", "Once every item is complete, the engineer submits it. This is the trigger that starts the approval chain."),
        ("Approvals sign-off chain", "The Migration Manager, Dev Lead, and (if needed) QA Lead review in a fixed order."),
        ("Delta Ready", "When all required approvals are in, the server is officially cleared to migrate."),
        ("Initiate Delta (Start)", "The engineer starts the real migration. The Manager is notified that it has begun."),
        ("Finish Delta", "When the migration completes, the engineer marks it finished and the Manager is notified."),
    ]
    for i, (t, s) in enumerate(steps_detail):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"Step {i+1} \u2014 {t}: ")
        rn.font.bold = True
        rn.font.color.rgb = RGBColor.from_string(_hex(PURPLE_DK))
        rn.font.size = Pt(11)
        r2 = p.add_run(s)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(_hex(INK))

    # ---- 7. Approval chain
    doc.add_page_break()
    add_heading(doc, "7. The approval chain explained")
    add_para(doc,
             "Approvals happen strictly in order \u2014 only the person whose turn it is can act. This keeps "
             "accountability clear and prevents anyone from jumping ahead.")
    add_image_centered(doc, appr, 5.8)
    for b in [
        "First the Migration Manager approves that the pre-checks are complete and valid.",
        "Next the Dev Lead approves \u2014 and decides whether a QA review is needed for this server.",
        "If QA is required, the QA Lead gives the final approval. If not, QA is skipped and the server becomes Delta Ready right away.",
        "If anyone declines, the chain moves back one step so the issue can be fixed and re-submitted \u2014 nothing is lost.",
        "Only when every required approval is in does the server become Delta Ready and the Start button unlocks.",
    ]:
        add_bullet(doc, b)

    # ---- 8. Tickets
    add_heading(doc, "8. Tickets \u2014 handling blockers")
    add_para(doc,
             "When something is blocking readiness, the team logs a ticket against the project and server, "
             "with a link to the full details. Tickets are marked Open or Resolved, so blockers are visible "
             "on the dashboard and never quietly forgotten. Only the engineer who created a ticket (or an "
             "Admin) can edit, resolve, or remove it.")

    # ---- 9. Notifications
    doc.add_page_break()
    add_heading(doc, "9. Staying informed \u2014 email notifications")
    add_para(doc,
             "The tool sends a clear, concise email each time the baton passes to the next person, so no one "
             "has to keep checking manually.")
    add_image_centered(doc, noti, 6.6)
    add_table(doc, ["When this happens", "Who receives an email"], [
        ["A pre-check is submitted", "Migration Manager \u2014 asked to review and approve"],
        ["The Manager approves", "Dev Lead \u2014 asked to review and approve"],
        ["The Dev Lead requires QA", "QA Lead \u2014 asked to review and approve"],
        ["All approvals are granted", "Migration Manager \u2014 notified the server is Delta Ready"],
        ["An engineer clicks Start", "Migration Manager \u2014 notified the Delta has been initiated"],
        ["An engineer clicks Finish", "Migration Manager \u2014 notified the Delta has finished"],
    ], widths=[2.9, 3.3])

    # ---- 10. Dashboard
    add_heading(doc, "10. The dashboard \u2014 everything at a glance")
    add_para(doc, "The home screen is designed to answer \u201cwhat needs attention right now?\u201d in seconds:")
    for b in [
        "Summary cards: total projects, servers, Delta-ready servers, pending approvals, open tickets, and projects ready to decommission.",
        "Three rings: server readiness, approvals approved-vs-pending, and overall project completion.",
        "Servers by product type: the mix of Content, Email, and Message work.",
        "Delta readiness by project: a progress bar per project, sorted so the most complete are on top.",
        "Needs Attention: a shortlist of projects with pending approvals or open tickets, each clickable to jump straight there.",
    ]:
        add_bullet(doc, b)

    # ---- 11. Status labels
    add_heading(doc, "11. Status labels you will see")
    add_table(doc, ["Label", "Meaning"], [
        ["Not started / Draft", "The pre-check has not been submitted yet."],
        ["Submitted / Pending", "Waiting for someone to review and approve."],
        ["Approved", "That role has signed off; the baton moves on."],
        ["Delta Ready", "All required gates are green \u2014 migration can begin."],
        ["Open (ticket)", "A blocker that still needs resolving."],
        ["Resolved (ticket)", "The blocker has been dealt with."],
    ], widths=[1.9, 4.3], header_bg=BLUE)

    # ---- 12. Example
    doc.add_page_break()
    add_heading(doc, "12. A day-in-the-life example")
    add_para(doc,
             "To see it all together: An engineer opens Project \u201cAcme Email Move,\u201d adds the server "
             "\u201cAcme-EX01,\u201d and imports 120 workspace pairs from a CSV. They work through the pre-check "
             "list \u2014 attaching a screenshot and a note to each item \u2014 then submit it.")
    add_para(doc,
             "The Migration Manager gets an email, opens Approvals, and approves. The Dev Lead is notified next, "
             "reviews, approves, and marks that QA is required. The QA Lead gets an email, checks the evidence, and "
             "gives the final approval. The server flips to Delta Ready and the Manager is notified.")
    add_para(doc,
             "The engineer clicks Start (the Manager is told the Delta has begun), runs the migration, and clicks "
             "Finish (the Manager is told it is done). Once every server in the project reaches this point, the "
             "project shows as ready to decommission.")

    # ---- 13. FAQ
    add_heading(doc, "13. Frequently asked questions")
    faqs = [
        ("Can a migration start before approvals are done?",
         "No. The Start button stays locked until the server is Delta Ready \u2014 that is the whole point of the tool."),
        ("What if I upload the same CSV twice?",
         "Identical rows are detected and skipped, and the skipped rows are shown to you in a popup. Only new rows are added."),
        ("What happens if an approver finds a problem?",
         "They decline, which sends the server back one step. The issue is fixed and re-submitted \u2014 no data or history is lost."),
        ("Is QA always required?",
         "No. The Dev Lead decides per server. If QA is not needed, it is skipped and the server becomes Delta Ready immediately."),
        ("Who can change or close a ticket?",
         "Only the engineer who created it, or an Admin."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rq = p.add_run("Q:  " + q)
        rq.font.bold = True
        rq.font.size = Pt(11)
        rq.font.color.rgb = RGBColor.from_string(_hex(PURPLE_DK))
        pa = doc.add_paragraph()
        pa.paragraph_format.space_after = Pt(8)
        ra = pa.add_run("A:  " + a)
        ra.font.size = Pt(11)
        ra.font.color.rgb = RGBColor.from_string(_hex(INK))

    # ---- 14. Summary
    add_divider(doc)
    add_heading(doc, "In one line")
    add_para(doc,
             "It turns \u201cis this server ready for Delta?\u201d from a question you chase across email into a "
             "single screen the whole team can trust \u2014 and the migration stays locked until every gate is green.",
             size=12.5, italic=True, color=PURPLE_DK)

    out = os.path.join(os.getcwd(), "Delta_Migration_Readiness_Tracker_Guide.docx")
    doc.save(out)
    print("SAVED:", out)


if __name__ == "__main__":
    build()
